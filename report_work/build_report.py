from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.numbering import NumberingPart
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\MobileApp")
WORK = ROOT / "report_work"
BASE = WORK / "official_base.docx"
OUTPUT = WORK / "ReEvent_Part2_draft.docx"
ASSETS = WORK / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)

BLACK = "000000"
MID_GREY = "E6E6E6"
PALE_GREY = "F2F2F2"
LIGHT_GREY = "F7F7F7"
DARK = "000000"
MUTED = "000000"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9029
SEQ_COUNTER = {"Figure": 0, "Table": 0}


def twips(inches: float) -> int:
    return int(round(inches * 1440))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def mark_non_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "false")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: Sequence[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BLACK, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size=12, bold=None, color=None, italic=None, name="Times New Roman") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_rpr_font(r_pr, name="Times New Roman") -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def clear_and_set(cell, text: str, *, bold=False, color=DARK, size=12, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_field(paragraph, instruction: str, result_text="") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    if result_text:
        result_run = paragraph.add_run(result_text)
        set_run_font(result_run, size=12)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, " PAGE ", "1")


def add_caption(container, label: str, title: str, *, above=False):
    SEQ_COUNTER[label] += 1
    p = container.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = above
    r = p.add_run(f"{label} ")
    set_run_font(r, 12, bold=True, color=BLACK)
    add_field(p, f" SEQ {label} \\* ARABIC ", str(SEQ_COUNTER[label]))
    r2 = p.add_run(f": {title}")
    set_run_font(r2, 12, color=DARK)
    return p


def set_picture_alt(paragraph, title: str, description: str) -> None:
    doc_props = paragraph._p.xpath(".//wp:docPr")
    if not doc_props:
        return
    doc_pr = doc_props[-1]
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, {1: 16, 2: 14, 3: 12}[level], bold=True, color=BLACK)
    return p


def add_back_heading(doc, text: str):
    p = doc.add_paragraph(text, style="Back Matter Heading")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, after=4, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_together = keep
    set_paragraph_spacing(p, after=after, line=1.5)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 12, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 12, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, 12, color=DARK)
    return p


def add_bullets(doc, items: Iterable[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, after=2, line=1.5)
        r = p.add_run(text)
        set_run_font(r, 12, color=DARK)


def add_callout(doc, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, BLACK, "8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREY)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=2, line=1.5)
    r = p.add_run(title + "  ")
    set_run_font(r, 12, bold=True, color=BLACK)
    r = p.add_run(text)
    set_run_font(r, 12, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_report_table(doc, title: str, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    add_caption(doc, "Table", title, above=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        set_cell_shading(hdr.cells[i], MID_GREY)
        clear_and_set(hdr.cells[i], text, bold=True, color=BLACK, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        for i, text in enumerate(row_data):
            if ridx % 2 == 1:
                set_cell_shading(row.cells[i], LIGHT_GREY)
            clear_and_set(row.cells[i], text, size=12)
    set_table_widths(table, widths)
    return table


def font(size: int, bold=False):
    candidates = [
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrapped(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, width: int, fnt, fill, spacing=7):
    words = text.split()
    lines, line = [], ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    draw.multiline_text((x, y), "\n".join(lines), font=fnt, fill=fill, spacing=spacing)
    return len(lines)


def make_architecture(path: Path) -> None:
    im = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(im)
    title_font, h_font, body_font, small_font = font(48, True), font(31, True), font(25), font(21)
    d.text((70, 42), "ReEvent implementation architecture", font=title_font, fill="#000000")
    d.text((70, 105), "Local-first Android experience with server-authoritative lifecycle operations", font=body_font, fill="#000000")

    def box(x, y, w, h, title, lines, fill, outline="#000000"):
        d.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=fill, outline=outline, width=4)
        d.text((x + 24, y + 18), title, font=h_font, fill="#000000")
        d.multiline_text((x + 24, y + 68), "\n".join(lines), font=small_font, fill="#000000", spacing=9)

    box(40, 185, 300, 250, "People and roles", ["Organiser", "Participant", "Partner"], "#F2F2F2")
    box(470, 165, 660, 305, "Android application", ["Jetpack Compose UI + role routing", "ViewModels and use cases", "Repository interfaces + domain policies", "Hilt dependency injection"], "#E6E6E6")
    box(1260, 185, 300, 250, "Public surfaces", ["Passport verifier website", "Partner/map tiles", "HTTPS QR resolver"], "#F2F2F2")
    box(250, 590, 430, 220, "On-device", ["Room scoped cache", "DataStore session/preferences", "WorkManager background sync"], "#F7F7F7")
    box(900, 560, 480, 250, "Supabase authority", ["Auth + row-level security", "Postgres + Storage", "RPC lifecycle commands", "Idempotent mutations"], "#F7F7F7")

    def arrow(x1, y1, x2, y2, label):
        d.line((x1, y1, x2, y2), fill="#000000", width=6)
        d.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill="#000000")
        d.text(((x1 + x2) // 2 - 70, (y1 + y2) // 2 - 30), label, font=small_font, fill="#000000")

    arrow(340, 310, 470, 310, "role-aware UI")
    arrow(1130, 310, 1260, 310, "public access")
    d.line((720, 470, 520, 590), fill="#000000", width=6)
    d.polygon([(520, 590), (526, 568), (540, 584)], fill="#000000")
    d.text((555, 510), "cache and sync", font=small_font, fill="#000000")
    d.line((900, 590, 820, 470), fill="#000000", width=6)
    d.polygon([(820, 470), (840, 480), (825, 493)], fill="#000000")
    d.text((835, 510), "typed gateways", font=small_font, fill="#000000")
    im.save(path, quality=95)


def make_placeholder(path: Path, title: str, capture: str) -> None:
    im = Image.new("RGB", (1200, 675), "#F7F7F7")
    d = ImageDraw.Draw(im)
    d.rounded_rectangle((28, 28, 1172, 647), radius=26, fill="#F7F7F7", outline="#000000", width=6)
    d.rectangle((28, 28, 1172, 118), fill="#E6E6E6", outline="#000000", width=2)
    d.text((60, 51), "SCREENSHOT PLACEHOLDER", font=font(36, True), fill="#000000")
    wrapped(d, title, 65, 175, 1070, font(46, True), "#000000", spacing=8)
    d.line((65, 320, 1135, 320), fill="#000000", width=3)
    d.text((65, 355), "Insert final ReEvent screenshot", font=font(31, True), fill="#000000")
    wrapped(d, capture, 65, 414, 1070, font(29), "#000000", spacing=10)
    im.save(path, quality=95)


def insert_figure(doc, image_path: Path, title: str, width_inches=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    set_picture_alt(p, title, "Diagram showing people and roles, the layered Android application, on-device storage, Supabase authority, and public services.")
    add_caption(doc, "Figure", title)


def add_placeholder_grid(doc, figures: Sequence[tuple[Path, str]], columns=2):
    rows = (len(figures) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [PAGE_WIDTH_DXA // columns] * columns
    set_table_widths(table, widths)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for idx, (path, title) in enumerate(figures):
        cell = table.cell(idx // columns, idx % columns)
        set_cell_shading(cell, "FFFFFF")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(2.8 if columns == 2 else 5.8))
        set_picture_alt(p, title, f"Placeholder for the final ReEvent screenshot showing {title}.")
        add_caption(cell, "Figure", title)
    if len(figures) % columns:
        cell = table.cell(rows - 1, columns - 1)
        cell.text = "Reserved for an additional screenshot if required."
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 12, italic=True, color=MUTED)
    return table


def add_section_page_number(section, fmt: str, start: int) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    footer.distance = Cm(1.1)
    p = footer.paragraphs[0]
    p.text = ""
    add_page_field(p)
    for run in p.runs:
        set_run_font(run, 12, color=MUTED)
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(8 if level == 1 else 5)
        style.paragraph_format.space_after = Pt(4 if level == 1 else 3)
        style.paragraph_format.keep_with_next = True
        set_rpr_font(style._element.get_or_add_rPr())

    # The official brief carries linked heading character styles that otherwise
    # reapply the Office Aptos theme when Word refreshes fields.
    for style_el in styles.element.findall(qn("w:style")):
        if style_el.get(qn("w:styleId")) in {
            "Heading1", "Heading2", "Heading3",
            "Heading1Char", "Heading2Char", "Heading3Char",
        }:
            r_pr = style_el.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                style_el.append(r_pr)
            set_rpr_font(r_pr)

    for style_name in ("List Bullet", "List Number"):
        if style_name not in [s.name for s in styles]:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        else:
            style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(2)

    for style_name in ("TOC 1", "TOC 2", "TOC 3", "Table of Figures"):
        if style_name in [s.name for s in styles]:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor.from_string(BLACK)

    if "Caption" not in [s.name for s in styles]:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.base_style = normal
    else:
        caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(DARK)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.line_spacing = 1.0

    if "Back Matter Heading" not in [s.name for s in styles]:
        style = styles.add_style("Back Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
    back = styles["Back Matter Heading"]
    back.font.name = "Times New Roman"
    back._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    back.font.size = Pt(16)
    back.font.bold = True
    back.font.color.rgb = RGBColor.from_string(BLACK)
    back.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    back.paragraph_format.line_spacing = 1.0
    back.paragraph_format.space_before = Pt(10)
    back.paragraph_format.space_after = Pt(6)
    back.paragraph_format.keep_with_next = True
    ppr = back._element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), "0")


def configure_heading_numbering(doc: Document) -> None:
    try:
        numbering_part = doc.part.numbering_part
    except NotImplementedError:
        numbering_element = OxmlElement("w:numbering")
        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            CT.WML_NUMBERING,
            numbering_element,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)
    numbering = numbering_part.element
    abstract_id = "42"
    num_id = "42"
    for el in list(numbering):
        if el.tag == qn("w:abstractNum") and el.get(qn("w:abstractNumId")) == abstract_id:
            numbering.remove(el)
        if el.tag == qn("w:num") and el.get(qn("w:numId")) == num_id:
            numbering.remove(el)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for ilvl, (pstyle, text) in enumerate((("Heading1", "%1"), ("Heading2", "%1.%2"), ("Heading3", "%1.%2.%3"))):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), pstyle)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + ilvl * 360))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + ilvl * 360))
        ind.set(qn("w:hanging"), "360")
        p_pr.extend([tabs, ind])
        r_pr = OxmlElement("w:rPr")
        set_rpr_font(r_pr)
        bold = OxmlElement("w:b")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), BLACK)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str({0: 32, 1: 28, 2: 24}[ilvl]))
        r_pr.extend([bold, color, size])
        lvl.extend([start, num_fmt, p_style, lvl_text, suff, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)
    for ilvl, style_name in enumerate(("Heading 1", "Heading 2", "Heading 3")):
        style = doc.styles[style_name]
        ppr = style._element.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            ppr.append(num_pr)
        for child in list(num_pr):
            num_pr.remove(child)
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), num_id)
        num_pr.extend([ilvl_el, num_id_el])


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_field_page(doc, heading: str, instruction: str, initial: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(heading)
    set_run_font(r, 16, bold=True, color=BLACK)
    holder = doc.add_paragraph()
    set_paragraph_spacing(holder, after=0)
    add_field(holder, instruction, initial)


def fill_cover(doc: Document) -> None:
    members = [
        ("LIEW KAIY BIN", "2302391"),
        ("MAH JUIN HONG", "TBC"),
        ("WONG JIE YING", "2401105"),
        ("WONG LOONG JIE", "2302662"),
    ]
    cover = doc.tables[0]
    for idx, (name, student_id) in enumerate(members, start=1):
        values = (name, student_id, "CS", "P3", "9", "ReEvent", "    /100")
        for col, value in enumerate(values):
            clear_and_set(cover.cell(idx, col), value, size=10 if col == 0 else 11, align=WD_ALIGN_PARAGRAPH.CENTER)
    names = doc.tables[2]
    for idx, (name, _) in enumerate(members, start=1):
        clear_and_set(names.cell(idx, 0), name, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    # The brief's logo is linked to an external URL and its embedded fallback is a
    # one-pixel strip. Replace only that broken fallback with the brief's own
    # embedded UTAR logo artwork extracted from the same source document.
    logo_paragraph = doc.paragraphs[3]
    for child in list(logo_paragraph._p):
        if child.tag != qn("w:pPr"):
            logo_paragraph._p.remove(child)
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.add_run().add_picture(str(WORK / "image3.jpeg"), width=Inches(2.15))
    set_picture_alt(logo_paragraph, "UTAR logo", "Universiti Tunku Abdul Rahman logo on the official assignment cover.")
    for table in doc.tables[:3]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    # Keep the official cover and marking tables unchanged apart from verified particulars.


def add_code_excerpt(doc: Document, path: Path, start: int, end: int) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()[start - 1:end]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, BLACK, "4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8F7")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.keep_together = False
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines, start=start):
        run = p.add_run(f"{idx:>3}  {line}" + ("\n" if idx < end else ""))
        set_run_font(run, 8.5, color=DARK, name="Times New Roman")


def build() -> None:
    make_architecture(ASSETS / "architecture.png")
    placeholders = [
        ("auth_role.png", "Authentication and role selection", "Capture sign-in or registration and the role-selection/routing state without real personal data."),
        ("organiser.png", "Organiser dashboard and event management", "Capture the organiser home/dashboard and one event-management screen using representative test data."),
        ("resource_editor.png", "Resource editor and persistent photo", "Capture create/edit resource with a persisted image, quantity, condition, and material classification."),
        ("passport_qr.png", "Digital passport and QR workflow", "Capture the passport plus a valid opaque HTTPS QR; do not expose account IDs or private test data."),
        ("marketplace.png", "Marketplace and transaction lifecycle", "Capture a listing and a representative request/handover/receipt state from the current implementation."),
        ("partner_matching.png", "Partner map, matching, and workbench", "Capture partner discovery and deterministic Smart/Circular Matching with an explanation."),
        ("impact.png", "Impact dashboard", "Capture repository-derived completed-outcome metrics; do not reuse the proposal's illustrative 83%, 1.8 t, or RM7.4k values."),
    ]
    for filename, title, capture in placeholders:
        make_placeholder(ASSETS / filename, title, capture)
    github_placeholder = (
        "github_commits.png",
        "GitHub commit history and contribution evidence",
        "Insert a GitHub commit history or contributor-insights screenshot showing the repository and all group member accounts.",
    )
    make_placeholder(ASSETS / github_placeholder[0], github_placeholder[1], github_placeholder[2])

    doc = Document(BASE)
    configure_styles(doc)
    configure_heading_numbering(doc)
    set_update_fields(doc)
    fill_cover(doc)

    # Apply A4 geometry while preserving the official cover's visible layout.
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Cm(1.1)
        section.footer_distance = Cm(1.1)

    # Front matter section: Roman numbering begins after the two-page official cover/marking sheets.
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    front.page_width = Cm(21.0)
    front.page_height = Cm(29.7)
    front.top_margin = Inches(1)
    front.bottom_margin = Inches(1)
    front.left_margin = Inches(1)
    front.right_margin = Inches(1)
    add_section_page_number(front, "lowerRoman", 1)

    add_field_page(doc, "Table of Contents", ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and update field.")
    doc.add_page_break()
    add_field_page(doc, "List of Figures", ' TOC \\h \\z \\c "Figure" ', "Right-click and update field.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("List of Tables")
    set_run_font(r, 16, bold=True, color=BLACK)
    holder = doc.add_paragraph()
    add_field(holder, ' TOC \\h \\z \\c "Table" ', "Right-click and update field.")

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    main.page_width = Cm(21.0)
    main.page_height = Cm(29.7)
    main.top_margin = Inches(1)
    main.bottom_margin = Inches(1)
    main.left_margin = Inches(1)
    main.right_margin = Inches(1)
    add_section_page_number(main, "decimal", 1)

    # Main body page 1.
    add_heading(doc, "Executive Summary and Introduction", 1)
    add_body(doc, "ReEvent is a role-based circular event resource platform designed to help organisers, participants, and recovery partners keep event materials in productive use. The implementation combines event and resource records, a circular marketplace, recovery programmes, digital passports, QR-assisted handovers, deterministic matching, and evidence-led impact reporting. It responds directly to Sustainable Development Goal 12 by encouraging responsible consumption, reuse, repair, redistribution, and recycling across an event's resource lifecycle (United Nations, n.d.).")
    add_heading(doc, "Problem and objectives", 2)
    add_body(doc, "Temporary events often create a fragmented trail of banners, furniture, decorations, packaging, equipment, and surplus stock. Their value is lost when ownership, condition, quantity, and recovery options are not visible at the moment a decision is made. ReEvent's objective is to make each item traceable, make circular routes discoverable, and make transaction completion auditable without exposing private identifiers in public QR payloads.")
    add_report_table(doc, "Assignment minimum-requirement compliance", ["Minimum requirement", "ReEvent evidence", "Status"], [
        ("Custom application icon", "ReEvent-specific launcher branding and Android resources are included in the application module.", "Completed"),
        ("Local storage", "Room provides account/environment-scoped records; DataStore retains session and preference state; WorkManager supports synchronisation.", "Completed"),
        ("Remote or public data source", "Supabase provides authentication, Postgres, Storage and RPC endpoints; the public website resolves passport URLs; MapTiler supplies map tiles.", "Completed"),
    ], [2050, 4860, 2119])
    # Main body page 2.
    section_heading = add_heading(doc, "Solution Evolution and Originality", 1)
    section_heading.paragraph_format.page_break_before = True
    add_body(doc, "The proposal established the circular-event problem and three-role experience. The implementation retains that identity while developing the prototype into a structured operational workflow. The result is an event-specific exchange rather than a generic resale application: resources remain connected to an origin event, a responsible custodian, a programme or transaction, and an auditable outcome.")
    add_report_table(doc, "Proposal-to-implementation traceability", ["Proposal direction", "Implemented evolution", "Reason for change"], [
        ("Landing concept", "Editorial role-aware entry and session restoration.", "Clarifies the next action and reduces navigation ambiguity."),
        ("Organiser, participant and partner roles", "Role routing and role-specific navigation/workspaces.", "Separates responsibilities and least-privilege actions."),
        ("Organiser recovery dashboard", "Event/resource management plus completed-outcome impact views.", "Connects management data to verifiable lifecycle events."),
        ("Circular marketplace", "Published listings with quantity-aware, server-authoritative requests.", "Prevents optimistic local state from becoming the source of truth."),
        ("Resource editor and photo capture", "Typed resource records and persisted photo metadata/storage.", "Improves traceability beyond a visual prototype."),
        ("Digital passport and QR", "Opaque-token HTTPS verifier with legacy-code safeguards.", "Avoids exposing resource or account identifiers."),
        ("AI recovery match", "Deterministic Smart/Circular Matching with explanations.", "Produces transparent and reproducible recommendations."),
        ("Partner map/network", "Partner discovery, programmes, geocoding controls and map presentation.", "Turns discovery into a programme request workflow."),
        ("Impact dashboard", "Metrics from completed transactions and documented factors only.", "Proposal values are illustrative, not evidence of achieved impact."),
    ], [2100, 4100, 2829])
    traceability_summary = add_body(doc, "Originality lies in the combination of event provenance, resource condition, circular-route matching, public passport verification, controlled custody transfers, and conservative impact visibility. These elements turn circular recovery into a traceable workflow that can be understood by organisers, participants, partners, and members of the public.")
    traceability_summary.paragraph_format.space_before = Pt(4)
    # Main body page 3.
    section_heading = add_heading(doc, "System Architecture and Technology", 1)
    section_heading.paragraph_format.page_break_before = True
    add_body(doc, "ReEvent uses a layered Kotlin architecture. Jetpack Compose renders adaptive role-based screens; ViewModels and use cases coordinate interaction; repository interfaces isolate storage and network details; and pure domain policies implement matching, QR validation, and impact summaries. Hilt supplies dependencies, Room caches account/environment-scoped state, DataStore persists lightweight preferences, and WorkManager coordinates background work.")
    insert_figure(doc, ASSETS / "architecture.png", "ReEvent system architecture", width_inches=6.0)
    add_heading(doc, "Remote, public and device services", 2)
    add_body(doc, "Supabase Auth, Postgres, Storage, row-level security, and typed RPC gateways provide the authoritative multi-user boundary. Critical lifecycle mutations include idempotency keys and are not completed offline. The public website resolves opaque passport tokens without revealing private record identifiers. MapLibre/MapTiler supports partner discovery, ZXing renders QR codes, and ML Kit supports scanning.")
    add_body(doc, "The local-first design improves responsiveness and read availability, while server authority protects shared quantities, permissions, custody, and completion states from conflicting device-side assumptions. Clear package boundaries, typed models, repository interfaces, and focused domain policies keep the implementation organised and make individual features easier to maintain.")
    # Main body pages 4-5.
    section_heading = add_heading(doc, "Implemented Functionalities and Completeness", 1)
    section_heading.paragraph_format.page_break_before = True
    add_heading(doc, "End-to-end user journeys", 2)
    add_body(doc, "An organiser can authenticate, create an event, record a resource and photo, publish it or route it to a programme, review requests, authorise a handover, and observe the completed outcome. A participant can discover a listing, request an eligible quantity, follow the handover/receipt or return path, and display an authorised passport. A partner can publish recovery programmes, appear in discovery, receive compatible requests, and operate the relevant workbench. Public users can resolve a valid passport URL through the verifier website.")
    add_report_table(doc, "Implemented functionality", ["Capability", "Implementation", "Status"], [
        ("Authentication and role routing", "Supabase-backed identity, session handling and role-specific destinations.", "Completed"),
        ("Events and resources", "Typed models, Room/cache boundaries, create/edit/status workflows and validation.", "Completed"),
        ("Resource photos", "Persisted local/remote metadata and Storage integration.", "Completed"),
        ("Marketplace", "Published listings, request quantities, eligibility and server-side lifecycle decisions.", "Completed"),
        ("Digital passports and QR", "Opaque HTTPS v1 payloads, validation and public verifier routing.", "Completed"),
        ("Partner programmes/map", "Programme authoring/discovery, map presentation and geocoding safeguards.", "Completed"),
        ("Smart/Circular Matching", "Deterministic condition/material/location ranking with explanations.", "Completed"),
        ("Impact dashboard", "Completed-transaction summaries and guarded numeric estimates.", "Completed"),
        ("Offline behaviour", "Scoped cache and read continuity; shared mutations remain online-only to protect consistency.", "Completed"),
        ("Account deletion", "Coordinated server and application deletion workflow.", "Completed"),
        ("Public passport website", "Public resolution of authorised passport links through opaque tokens.", "Completed"),
    ], [1950, 4740, 2339])
    add_heading(doc, "Lifecycle integrity and data boundaries", 2)
    add_body(doc, "A lifecycle command gateway turns user intent into typed Supabase RPC calls. The server validates actors, resource side, quantities, open listings, category/condition acceptance, custody, and state transitions. Idempotency keys make retries safe and prevent the same client intention from being applied as a different operation. Failure codes are classified so the UI can distinguish retryable transport problems from terminal rule violations.")
    add_heading(doc, "Matching and impact claims", 2)
    add_body(doc, "Smart/Circular Matching uses a deterministic rules engine. It rejects inactive or empty resources, orders routes by condition, filters active programmes by material, considers location compatibility, and returns an explanation for the recommended route. The transparent result helps users understand why reuse, donation, repair, resale, or recycling is suggested.")
    add_body(doc, "Impact totals are derived from completed transactions and valid impact records. The only current numeric emissions estimate is a demonstration for completed plastic/acrylic recycling recorded in kilograms, using a documented factor of 1.59710826 kgCO2e per kilogram. The factor is derived from UK Government 2025 conversion-factor material and is presented as a conservative product estimate, not a full life-cycle assessment (Department for Energy Security and Net Zero, 2025). When evidence is absent, the application displays an unavailable reason instead of inventing a value.")
    # Main body pages 6-7.
    section_heading = add_heading(doc, "UI/UX Design and Creativity", 1)
    section_heading.paragraph_format.page_break_before = True
    add_body(doc, "The implementation evolves the proposal into an editorial, task-focused interface. Role-specific navigation reduces irrelevant choices; prominent states explain loading, emptiness, connectivity, permission, and validation outcomes; and adaptive Compose layouts support different phone dimensions. Dark green, neutral surfaces, strong hierarchy, and restrained cards preserve the sustainability identity without allowing decoration to obscure workflow state.")
    add_heading(doc, "Final UI evidence placeholders", 2)
    add_body(doc, "The reserved figure areas below are sized for screenshots from the final ReEvent application. The existing captions are connected to the automatic List of Figures and can be retained when each image is inserted.")
    add_placeholder_grid(doc, [(ASSETS / f, t) for f, t, _ in placeholders[:4]], columns=2)
    section_heading = add_heading(doc, "Role workflows and visual continuity", 2)
    section_heading.paragraph_format.page_break_before = True
    add_body(doc, "The organiser's view emphasises event and recovery management; the participant's view emphasises discovery, requests, passports and returns; and the partner's view emphasises programmes, compatible resources and operational decisions. The following placeholders cover the remaining recovery, matching and impact surfaces.")
    add_placeholder_grid(doc, [(ASSETS / f, t) for f, t, _ in placeholders[4:]], columns=2)
    add_body(doc, "The placeholders preserve consistent image proportions and caption positions across the seven key interfaces.")
    # Main body page 8.
    section_heading = add_heading(doc, "Team Contributions", 1)
    section_heading.paragraph_format.page_break_before = True
    add_body(doc, "Development responsibilities were divided by feature area and integrated through the shared ReEvent repository. Each member contributed to a distinct part of the implemented workflow while supporting integration and review across the application.")
    add_report_table(doc, "Member roles and contributions", ["Member", "Primary contribution areas"], [
        ("Liew Kaiy Bin (2302391)", "Core platform, authentication, persistence, Supabase integration, and application integration."),
        ("Mah Juin Hong", "Marketplace, listing interactions, and partner programme workflows."),
        ("Wong Jie Ying (2401105)", "Event and resource lifecycle, interface implementation, photos, passports, and QR workflow."),
        ("Wong Loong Jie (2302662)", "Circular matching, impact reporting, quality review, and deployment preparation."),
    ], [2700, 6329])
    add_heading(doc, "GitHub contribution evidence", 2)
    add_body(doc, "The figure below is reserved for a GitHub commit-history or contributor-insights screenshot that shows the repository and the contribution record for all group members.")
    add_placeholder_grid(doc, [(ASSETS / github_placeholder[0], github_placeholder[1])], columns=1)

    # Main body page 9.
    final_heading = add_heading(doc, "Civil and Commercial Potential", 1)
    final_heading.paragraph_format.page_break_before = True
    add_body(doc, "ReEvent can support universities, municipal councils, event organisers, venues, NGOs, reuse networks, repairers and recycling partners. A campus could retain ownership and condition history for reusable staging assets; a venue could expose surplus resources after an exhibition; and a recovery partner could publish accepted materials, service areas and programme rules. Commercial pathways include organisation workspaces, managed partner networks, compliance exports, verified material passports, and service fees for coordinated recovery, while civil deployments may prioritise open access and community redistribution.")
    add_heading(doc, "Limitations and risks", 2)
    add_bullets(doc, [
        "Shared lifecycle actions and public passport verification require network access because authoritative quantities, permissions, and custody changes are held by the server.",
        "Map and geocoding functions depend on provider availability, valid configuration, usage quotas, attribution, and service terms.",
        "Impact figures depend on accurate resource quantities, compatible measurement units, completed transaction records, and documented conversion factors.",
        "Public passport access requires controlled token issuance, revocation, privacy review, and operational monitoring.",
        "The value of programme matching grows with the number and coverage of participating recovery partners.",
    ])
    add_heading(doc, "Future enhancements", 2)
    add_body(doc, "Future releases can extend organisation-level dashboards, downloadable sustainability reports, multilingual content, accessibility options, partner onboarding tools, and additional material conversion factors. With sufficient consented historical data, the existing transparent matching rules could also support an optional learned ranking layer while retaining explanations and a deterministic fallback. Wider deployment can include production monitoring and additional mobile distribution channels.")
    add_heading(doc, "Conclusion", 2)
    add_body(doc, "ReEvent delivers a complete circular event-resource workflow with role-aware journeys, traceable resources, authoritative lifecycle commands, privacy-conscious passports, deterministic recovery matching, and evidence-led impact reporting. By connecting organisers, participants, and recovery partners around the same resource history, the application converts the proposal's circular-event concept into a practical platform that supports responsible consumption and recovery decisions.")

    # References are outside the under-ten-page substantive body.
    references_heading = add_back_heading(doc, "References")
    references_heading.paragraph_format.page_break_before = True
    refs = [
        "Department for Energy Security and Net Zero. (2025). Greenhouse gas reporting: Conversion factors 2025. UK Government. https://www.gov.uk/government/publications/greenhouse-gas-reporting-conversion-factors-2025",
        "ReEvent Group. (2026). ReEvent proposal [Unpublished group assignment proposal]. Universiti Tunku Abdul Rahman.",
        "United Nations. (n.d.). Goal 12: Ensure sustainable consumption and production patterns. Sustainable Development Goals. https://sdgs.un.org/goals/goal12",
    ]
    for ref in refs:
        p = add_body(doc, ref, after=6)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)

    # Appendices.
    appendix_specs = [
        ("Appendix A — Privacy-conscious QR payload contract", ROOT / r"ReEvent\app\src\main\java\com\reevent\app\feature\passports\PassportQrPayload.kt", 47, 79,
         "Creates and validates the canonical HTTPS verifier URL while rejecting malformed or unconfigured values.", "Wong Jie Ying"),
        ("Appendix B — Circular Matching Excerpt", ROOT / r"ReEvent\app\src\main\java\com\reevent\app\feature\matching\CircularRecommendationEngine.kt", 11, 33,
         "Ranks eligible recovery routes by resource state, material compatibility and location, returning an explanation rather than an opaque AI score.", "Wong Loong Jie"),
        ("Appendix C — Impact calculation guardrails", ROOT / r"ReEvent\app\src\main\java\com\reevent\app\feature\impact\ImpactCalculator.kt", 11, 42,
         "Uses completed transactions and valid impact records, and reports why estimates are unavailable.", "Wong Loong Jie"),
        ("Appendix D — Server-authoritative lifecycle gateway", ROOT / r"ReEvent\app\src\main\java\com\reevent\app\core\network\LifecycleCommandGateway.kt", 46, 87,
         "Converts typed commands into idempotent Supabase RPC requests so shared lifecycle state is decided by the server.", "Liew Kaiy Bin"),
    ]
    for title, path, start, end, purpose, contributor in appendix_specs:
        appendix_heading = add_back_heading(doc, title)
        appendix_heading.paragraph_format.page_break_before = True
        add_report_table(doc, f"{title} metadata", ["File path", "Purpose", "Contributor"], [(str(path.relative_to(ROOT)), purpose, contributor)], [3000, 4200, 1829])
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(1)
        set_run_font(spacer.add_run(" "), 1, color=BLACK)
        add_code_excerpt(doc, path, start, end)
        add_body(doc, "Excerpt note: This is a focused implementation excerpt rather than a complete source listing. The repository version remains authoritative.", after=0)

    # Final global font normalisation.
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(BLACK)
            if p.style and p.style.name.startswith("Heading"):
                continue
            if p.style and p.style.name == "Back Matter Heading":
                continue
            if run.font.size is None:
                set_run_font(run, 12, color=DARK)
    for table in doc.tables[3:]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(BLACK)
                        if run.font.size is None:
                            set_run_font(run, 12, color=DARK)

    # Explicitly mark non-header layout/code tables without asking Word to repeat
    # them, while retaining real repeating header rows on data tables.
    for table in doc.tables:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            mark_non_header_row(table.rows[0])

    doc.core_properties.title = "ReEvent Part 2 Implementation Report"
    doc.core_properties.subject = "UCCD3223 Mobile Applications Development — Group Assignment 2"
    doc.core_properties.author = "Liew Kaiy Bin; Mah Juin Hong; Wong Jie Ying; Wong Loong Jie"
    doc.core_properties.keywords = "ReEvent, circular economy, mobile application, implementation report"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
