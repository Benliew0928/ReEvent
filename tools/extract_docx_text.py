from pathlib import Path
import sys

from docx import Document


def iter_block_text(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table_index, table in enumerate(doc.tables, start=1):
        yield f"[Table {table_index}]"
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                yield " | ".join(cells)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: extract_docx_text.py input.docx output.txt")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(str(input_path))
    output_path.write_text("\n".join(iter_block_text(doc)), encoding="utf-8")


if __name__ == "__main__":
    main()
