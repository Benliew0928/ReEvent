from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ReEvent_Mobile_App_Proposal.docx"
ASSETS = ROOT / "proposal_assets"
ASSETS.mkdir(exist_ok=True)


INK = RGBColor(16, 24, 32)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(80, 88, 96)
GREEN = RGBColor(14, 124, 102)


def style_run(run, bold=False, color=None, size=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_para(doc, text="", align=None, size=11, color=INK, space_after=8, line_spacing=1.25):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    r = p.add_run(text)
    style_run(r, color=color, size=size)
    return p


def add_caption(doc, text):
    p = add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=DARK_BLUE, space_after=4, line_spacing=1.0)
    for run in p.runs:
        run.bold = True
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        style_run(run, bold=True, color=BLUE if level <= 2 else DARK_BLUE, size=16 if level == 1 else 13)
    return p


def create_architecture_diagram(path):
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 48)
        box_font = ImageFont.truetype("arialbd.ttf", 34)
        small_font = ImageFont.truetype("arial.ttf", 26)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()

    draw.text((70, 45), "ReEvent Overall System Architecture", fill="#101820", font=title_font)

    boxes = {
        "Mobile App": (80, 170, 430, 390, "#F8F2E8", ["Jetpack Compose UI", "Role dashboards", "QR scan and item forms"]),
        "App Logic": (550, 170, 930, 390, "#EAF4F0", ["MVVM view models", "Circular pathway rules", "Impact calculations"]),
        "Local Storage": (1050, 170, 1430, 390, "#EEF2FF", ["Room database", "Offline resource drafts", "Scan history cache"]),
        "Cloud Backend": (550, 530, 930, 750, "#F4F6F9", ["Supabase or Firebase", "Auth, events, resources", "Partner programmes"]),
        "External Services": (1050, 530, 1430, 750, "#FFF4DF", ["Maps API", "AI image assessment", "Notifications"]),
        "Users": (80, 530, 430, 750, "#FDECEF", ["Organisers", "Participants", "Circular partners"]),
    }

    for title, (x1, y1, x2, y2, fill, lines) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline="#AAB7C2", width=4)
        draw.text((x1 + 26, y1 + 24), title, fill="#0E7C66", font=box_font)
        y = y1 + 88
        for line in lines:
            draw.text((x1 + 32, y), line, fill="#2F3A3D", font=small_font)
            y += 42

    arrows = [
        ((430, 280), (550, 280)),
        ((930, 280), (1050, 280)),
        ((740, 390), (740, 530)),
        ((930, 640), (1050, 640)),
        ((430, 640), (550, 640)),
        ((260, 390), (260, 530)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#2563EB", width=8)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            pts = [(ex, ey), (ex - 22, ey - 14), (ex - 22, ey + 14)]
        elif ey > sy:
            pts = [(ex, ey), (ex - 14, ey - 22), (ex + 14, ey - 22)]
        else:
            pts = [(ex, ey), (ex + 22, ey - 14), (ex + 22, ey + 14)]
        draw.polygon(pts, fill="#2563EB")

    draw.text(
        (80, 850),
        "Key flow: scan resource -> create digital passport -> recommend circular pathway -> match partner/user -> record recovery impact",
        fill="#101820",
        font=small_font,
    )
    img.save(path)


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    p = add_para(doc, "UNIVERSITI TUNKU ABDUL RAHMAN", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, space_after=18)
    p.runs[0].bold = True

    add_para(doc, "Mobile Applications Development\nGroup Assignment Part 1: Proposal", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=MUTED, space_after=34)

    p = add_para(doc, "ReEvent", align=WD_ALIGN_PARAGRAPH.CENTER, size=30, color=GREEN, space_after=8)
    p.runs[0].bold = True

    p = add_para(doc, "Circular Event Resource Management Mobile Application", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=INK, space_after=24)
    p.runs[0].bold = True

    logo = ROOT / "reevent_logo_generated" / "reevent-logo-concepts-02-app-icon-sheet.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.add_run().add_picture(str(logo), width=Inches(3.6))

    add_para(doc, "Prepared by: [Group member names]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    add_para(doc, "Programme / Course: [Programme name] / Mobile Applications Development", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    add_para(doc, "Lecturer / Tutor: [Lecturer or tutor name]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    add_para(doc, "Submission date: [Submission date]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=4)
    add_para(doc, "Project focus: SDG 12, Responsible Consumption and Production", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=8)

    doc.add_page_break()


def add_image(doc, path, width, caption=None):
    if caption:
        add_caption(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(path), width=Inches(width))


def add_ui_mockups(doc):
    screenshots = [
        (
            "Onboarding / Role Select and Organiser Command Centre mock-ups",
            ROOT / "figma_screenshots" / "onboarding.png",
            ROOT / "figma_screenshots" / "home.png",
        ),
        (
            "AI Circular Match and Event Impact Dashboard mock-ups",
            ROOT / "figma_screenshots" / "ai_scan.png",
            ROOT / "figma_screenshots" / "impact.png",
        ),
    ]
    for caption, left, right in screenshots:
        if left.exists() and right.exists():
            add_caption(doc, caption)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            p.add_run().add_picture(str(left), width=Inches(2.2))
            p.add_run("     ")
            p.add_run().add_picture(str(right), width=Inches(2.2))


def build():
    create_architecture_diagram(ASSETS / "architecture.png")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    add_cover(doc)

    add_heading(doc, "1. Introduction", 1)
    add_para(doc, "Modern events depend on many short-term resources, including booth panels, banners, badge holders, decorations, reusable cups, table covers, signage, extension cables, props, display stands, light equipment, packaging materials, and gift bags. These items are often purchased or prepared for a single event, used for a short period, and then left without a clear plan after the event ends. Some items may be thrown away even though they are still usable, while others may be stored in offices, club rooms, or storerooms until they are forgotten. This situation creates unnecessary spending for organisers and increases avoidable waste from event activities.")
    add_para(doc, "ReEvent is proposed as a mobile application that helps event organisers, participants, and circular partners keep event resources in circulation for as long as possible. The project is connected to the circular economy idea, where products and materials should be reused, repaired, shared, returned, refurbished, or recycled instead of becoming waste after one short use. It also supports Sustainable Development Goal 12, Responsible Consumption and Production, because the application encourages users to manage resources more responsibly and make better decisions before, during, and after an event.")
    add_para(doc, "The first target environment for ReEvent is university and community events because these events are common, familiar, and realistic for a student mobile application project. A university career fair, student club exhibition, open day, charity bazaar, school programme, or community fair usually involves many reusable items, yet the recovery process is often informal. By focusing on this environment, the proposal can show a clear problem, a practical user group, and a solution that can be demonstrated through a polished mobile prototype. In the future, the same concept can be expanded to larger conferences, exhibitions, public events, and professional event agencies.")
    add_para(doc, "The proposal presents ReEvent as more than a simple recycling or second-hand marketplace app. The intention is to design a mobile system that helps users plan the resource lifecycle from the beginning of the event, not only after waste has already been created. Through digital resource passports, QR scanning, circular pathway recommendations, partner matching, and impact tracking, ReEvent aims to make sustainable event management easier, more organised, and more measurable.")

    add_heading(doc, "2. Problem Statement, Motivation, and Market Gap Analysis", 1)
    add_heading(doc, "2.1 Problem Statement", 2)
    add_para(doc, "The main problem addressed by ReEvent is that event waste is not only a disposal issue, but also a resource coordination issue. Event organisers may not have a complete record of what materials were purchased, borrowed, rented, or prepared for an event. After the event, they may not know which items are still usable, which items need repair, which items can be donated, and which items should be sent to a suitable recovery partner. Without a simple system to manage this information, resources that still have value may be treated as rubbish or left unused.")
    add_para(doc, "Participants and community users also face a similar problem because they may not know what can be returned, reused, exchanged, or claimed after an event. For example, badge holders, reusable cups, decorations, small props, and event merchandise may have a possible next use, but participants often do not receive clear instructions about what to do with them. Circular partners, such as repair shops, recyclers, refurbishers, rental suppliers, donation centres, and second-hand resellers, may also be disconnected from event organisers because they do not receive enough details about material type, quantity, condition, and location. This missing connection causes useful materials to lose value and makes sustainable recovery difficult.")
    add_para(doc, "Another part of the problem is that many event teams do not have measurable evidence of their sustainability efforts. Even when an organiser tries to reuse or donate materials, the result is often not recorded in a structured way. This makes it difficult to show how many items were recovered, how much purchasing was avoided, or how much waste was prevented. As a result, sustainability becomes an informal intention rather than a clear and trackable event management process.")
    add_heading(doc, "2.2 Motivation", 2)
    add_para(doc, "The motivation for developing ReEvent comes from the observation that many events create temporary consumption even when the materials involved are still useful after the event ends. University events, exhibitions, fairs, and community programmes often require repeated types of resources, but each organiser may still buy or prepare new materials because they do not know what is already available. This creates a cycle where event materials are used for a short time, forgotten, and then replaced by new purchases for the next event.")
    add_para(doc, "ReEvent is motivated by the need to help organisers think beyond the event day. A good event should not only focus on attendance, decoration, and smooth operation, but also on what happens to the materials after the event is completed. If organisers can scan resources, update their condition, and choose a next action immediately after the event, then reuse and recovery become part of the normal workflow instead of an extra task that is easily ignored.")
    add_para(doc, "The project is also motivated by the opportunity to encourage behavioural change. Participants may be more willing to return reusable items, exchange resources, or support sustainable actions when the process is visible and convenient through a mobile app. Organisers may also become more aware of their consumption patterns when they can see the number of items recovered, reused, repaired, or sent to partners. In this way, ReEvent does not only provide a technical function, but also encourages a more responsible habit in event planning.")
    add_heading(doc, "2.3 Market Gap Analysis", 2)
    add_para(doc, "Current event management applications usually focus on schedules, ticketing, registration, check-in, attendee communication, and event promotion. These functions are important for running an event, but they normally stop at the point where the event is successfully delivered. They do not usually provide a complete post-event resource recovery workflow that tracks what happens to booth materials, signs, props, reusable items, and leftover assets after the event has ended.")
    add_para(doc, "General marketplace applications allow users to buy and sell second-hand items, but they are not designed around event resource management. They do not normally store event ownership, reuse history, repair history, material category, recovery pathway, or sustainability impact. Recycling locator applications can help users find recycling points, but they usually treat items as waste after use instead of asking whether the item can be reused, repaired, shared, rented, donated, refurbished, or returned before recycling becomes necessary.")
    add_para(doc, "This leaves a market gap for a focused mobile solution that connects event operations with circular economy practices. Many event teams still use spreadsheets, messaging apps, paper notes, and memory to manage materials, which makes the process fragmented and difficult to track. ReEvent fills this gap by giving event resources a digital identity and by connecting organisers, participants, and circular partners in one system. The app is positioned between event management, circular marketplace, and sustainability reporting, which makes it different from apps that only handle one of these areas.")

    add_heading(doc, "3. Proposed Idea, Solution, and Contribution", 1)
    add_heading(doc, "3.1 Proposed Idea", 2)
    add_para(doc, "The proposed idea is ReEvent, a circular event resource management mobile application. The app helps users register event resources before an event, scan and update items during the event, and choose the most suitable recovery action after the event. Instead of viewing event materials as temporary objects that are only useful for one occasion, ReEvent encourages users to treat each resource as something with a lifecycle, a condition, a history, and a possible next user.")
    add_para(doc, "The idea is built around a simple event scenario. Before an event begins, the organiser creates an event profile and registers important resources such as booth panels, badge holders, reusable cups, acrylic signs, table covers, display stands, storage boxes, cables, and gift bags. During the event, the organiser or volunteers can scan the items to confirm their status. After the event, ReEvent helps the organiser decide whether each resource should be kept for a future event, shared with another organiser, rented, sold, donated, repaired, refurbished, returned to a supplier, or recycled through an approved partner.")
    add_para(doc, "The app is intended to support three main user groups. Event organisers use the app to manage resources, reduce unnecessary purchasing, and report recovery outcomes. Participants and community users use the app to return items, claim reusable materials, exchange event resources, or take part in sustainable actions. Circular partners use the app to receive suitable recovery requests and state what kinds of materials, conditions, and quantities they can accept.")
    add_heading(doc, "3.2 Solution", 2)
    add_para(doc, "The main solution in ReEvent is the Digital Resource Passport. Every item registered under ReEvent will be assigned a QR-linked record that stores information such as item name, category, material, owner, quantity, location, condition, reuse count, estimated value, and recommended circular pathway. The purpose of the QR code is to allow users to obtain the resource details after scanning it and guide them towards the next suitable action, such as viewing the payment method, pickup location, return option, or recovery pathway. This makes all registered items traceable because the physical object is connected to a clear digital identity. We chose this solution because many items lose their value when their details are separated from the physical item. For example, a stack of booth panels may still be usable, but if no one knows who owns them, what material they are made from, or where they are stored, the panels may be forgotten and may not be reused elsewhere. ReEvent's QR-linked passport solves this problem because all registered item details are recorded, and users only need to scan the QR code of the specific item to view the necessary information.")
    add_para(doc, "Besides that, ReEvent also uses circular pathway recommendation to help users decide what should happen to each item. The app will prioritise reuse before recycling because reuse usually keeps the item at a higher value. For example, a booth panel that is in good condition can be stored for the next event or rented out to another organiser. A damaged but repairable panel can be sent to a repair partner and reused for a later event. An unwanted but usable panel can be donated to a school or community group. Only when an item has no reuse, repair, or refurbishment option will it be recommended for recycling. The recommendation process is based on several factors, such as item category, condition, material type, and current location.")
    add_para(doc, "This solution is helpful in ReEvent because it connects the app's main functions into one clear workflow. The Digital Resource Passport gives each item an identity, while the circular pathway recommendation helps users decide the next action based on that identity. Organisers can manage their resources more systematically, participants can understand what they should do with event items, and circular partners can receive clearer information about the resources they may accept. As a result, ReEvent does not only record items, but also helps move them towards reuse, repair, donation, rental, or recycling in a more organised way.")
    add_heading(doc, "3.3 Contribution", 2)
    add_para(doc, "The main contribution of ReEvent is that it combines event resource planning, circular economy action, and sustainability reporting in one mobile application. It is not only a recycling app because it considers reuse, repair, sharing, donation, supplier return, refurbishment, and partner matching before recycling. It is also not only a marketplace because it tracks item identity, event history, condition, ownership, and impact. ReEvent contributes to organisers by reducing unnecessary purchases, helps participants return or claim reusable resources more conveniently, and gives circular partners clearer information about recoverable items. From a mobile application development perspective, it also demonstrates useful features such as user roles, local storage, QR scanning, resource forms, dashboards, map-based partner discovery, and backend synchronisation.")

    add_heading(doc, "4. Initial Design of Solution", 1)
    add_heading(doc, "4.1 Overall System Architecture Design", 2)
    add_para(doc, "The proposed system uses a mobile-first architecture. The Android mobile application provides the main user experience through role-based screens for event organisers, participants, and circular partners. The app logic layer manages user actions, resource status changes, circular pathway recommendations, and impact calculations. Local storage supports offline event work by saving resource drafts, scan history, cached event data, and partner programme information. The cloud backend synchronises user accounts, event records, resource passports, partner programmes, and recovery outcomes across devices.")
    add_image(doc, ASSETS / "architecture.png", width=6.35, caption="Overall system architecture")
    add_para(doc, "The architecture begins when a user creates or joins an event in the mobile app. An organiser can register a resource by entering item details and attaching a photo. The system then creates a digital resource passport and connects it to a QR code. When the QR code is scanned later, the app retrieves the passport, checks the current condition, and allows the user to update the status. This flow allows the same item to be tracked across different event stages instead of being recorded only once.")
    add_para(doc, "The cloud backend is responsible for storing shared information that needs to be accessed by multiple users, such as event details, resource records, partner programmes, marketplace listings, and recovery outcomes. Local storage is still important because event venues may have unstable internet access, especially during setup or teardown. By using local storage, the app can save scan history and draft updates first, then synchronise with the backend when the connection becomes available.")
    add_para(doc, "External services can improve the usefulness of the application. A map service can help users locate nearby circular partners, donation centres, repair shops, or recycling points. An optional AI image assessment service can support item categorisation or condition suggestion based on a photo, although the final decision should still be confirmed by the user. Notification services can remind organisers to complete post-event recovery tasks or remind participants to return deposit-based items before a deadline.")
    add_heading(doc, "4.2 Mock-up UI Design of Mobile Application", 2)
    add_para(doc, "The mobile interface is designed to feel like a professional event operations and sustainability tool rather than a generic recycling application. The visual direction uses clean dashboards, compact resource cards, QR passport elements, partner map screens, progress indicators, and realistic event-resource data. The app should look practical and trustworthy because the target users include organisers who need to manage real materials, not only casual users browsing environmental tips.")
    add_para(doc, "The proposed navigation begins with onboarding and role selection. In this screen, the user chooses whether they are an event organiser, participant, or circular partner, which allows the app to show a more relevant workflow. After signing in, an organiser enters the command centre, where the current event, recovery progress, registered resources, and urgent actions are displayed. This screen acts as the main control point because it allows the organiser to scan an item, add a new resource, check pending recovery actions, or move to the impact dashboard.")
    add_para(doc, "The Digital Resource Passport screen is one of the most important screens because it shows the identity and status of an item. It should display the item name, image, QR code, category, material, condition, owner, event history, and recommended next action. The AI Circular Match screen can support the user by suggesting whether an item should be reused, repaired, donated, returned, or recycled based on its category and condition. These screens help make the circular economy process understandable because the user can see the reason behind each recommendation.")
    add_para(doc, "The marketplace and partner discovery screens extend the app beyond simple inventory management. Through the marketplace, organisers and participants can share, rent, sell, exchange, or donate suitable resources. Through the partner map, users can find circular partners who accept specific materials or services, such as repair, refurbishment, take-back, donation, or recycling. The impact dashboard completes the experience by showing the number of recovered items, reuse progress, estimated waste avoided, and estimated cost saved. Together, these screens show how the app helps users move from planning to action and finally to measurable impact.")
    add_ui_mockups(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build()
